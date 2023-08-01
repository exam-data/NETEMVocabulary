import docx
import json

def split_string_with_commas(input_str):
    result = []
    current_word = ""
    for char in input_str:
        if char == "、":  
            if current_word:
                result.append(current_word + char)  
                current_word = ""
        else:
            current_word += char
    if current_word:
        result.append(current_word)
    return result

def concatenate_with_newlines(str_list):
    result = str_list[0]
    for i in range(1, len(str_list)):
        current_length = len(result)
        next_length = len(str_list[i])
        if current_length + next_length > 6 and current_length < 6:
            result += '\n'
        result += str_list[i]
    return result

def format_word(word):
    single_words = split_string_with_commas(word)
    return concatenate_with_newlines(single_words)

# 读取 JSON 文件内容
with open("vocabulary.json", "r", encoding="utf-8") as file:
    data = json.load(file)

doc = docx.Document('5530_v4.1.docx')

# 创建一个列表，用于存储未找到的单词
not_found_words = []

# 遍历文档中的所有表格
for table in doc.tables:
    # 假设单词列是第三列（从0开始数）
    for row in table.rows:
        # 跳过表头行
        if row.cells[0].text == '序号':
            continue
        # 获取单词列的单元格内容
        word = row.cells[2].text
        # 判断单词是否为空
        if word.strip() != "":
            word_found = False
            for entry in data["5530考研词汇词频排序表"]:
                if entry["单词"] == word:
                    definition = entry["释义"]
                    # 格式化释义
                    definition = format_word(definition)
                    # 将 JSON 数据中的释义写入单词所在行的第四列
                    row.cells[3].text = definition
                    word_found = True
                    break
            if not word_found:
                not_found_words.append(word)

# 输出未找到的单词
if not_found_words:
    print("未找到以下单词：")
    for word in not_found_words:
        print(word)

# 设置所有表格中的文字居中对齐
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

# 保存更新后的 Word 文档
doc.save("updated_5530_v4.1.docx")

import docx

def split_string_with_commas(input_str):
    result = []
    current_word = ""
    for char in input_str:
        if char == "、":  
            if current_word:
                result.append(current_word + char)  
                current_word = ""
        else:
            current_word += char
    if current_word:
        result.append(current_word)
    return result

def concatenate_with_newlines(str_list):
    result = str_list[0]
    for i in range(1, len(str_list)):
        current_length = len(result)
        next_length = len(str_list[i])
        if current_length + next_length >= 6 and current_length <= 6:
            result += '\n'
        result += str_list[i]
    return result

def format_word(word):
    single_words = split_string_with_commas(word)
    return concatenate_with_newlines(single_words)

def remove_whitespace(input_str):
    cleaned_str = input_str.replace(" ", "").replace("\t", "").replace("\n", "").replace("\r", "")
    return cleaned_str

doc = docx.Document('5530考研词汇词频排序表12.docx')

for table in doc.tables:
   
    for row in table.rows:
        # 跳过表头行
        if row.cells[0].text == '序号':
            continue
        # 获取单词列的单元格内容
        definition = row.cells[4].text
        if definition:
            definition=remove_whitespace(definition)
            row.cells[4].text = format_word(definition)

# 设置所有表格中的文字居中对齐
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

# 保存更新后的 Word 文档
doc.save("updated_5530_v4.9.docx")
